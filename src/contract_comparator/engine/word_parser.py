"""
Word 文档解析模块
使用 python-docx 提取 Word 文档的文本和结构化信息
"""
import os
import re
from lxml import etree
from docx import Document
from docx.shared import Pt


class WordParser:
    """Word 文档解析器"""
    
    def __init__(self, docx_path: str):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Word 文件不存在: {docx_path}")
        
        self.docx_path = docx_path
        self.document = Document(docx_path)
        self.full_text = ""
        self.paragraphs = []
        self.tables = []
    
    @staticmethod
    def extract_text_with_revisions(para_element) -> str:
        """
        从段落 XML 中提取文本，处理修订标记（插入和删除）

        - <w:ins> 插入的内容：保留
        - <w:del> 删除的内容：移除（不在最终文本中出现）

        python-docx 默认忽略修订标记，此方法手动解析 XML。
        """
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        }

        text_parts = []

        # 遍历所有 <w:t> 文本节点
        for t_elem in para_element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            # 检查这个 <w:t> 是否位于 <w:del> 删除标记内
            # 如果是，说明这段文本已被删除，不应包含在结果中
            parent = t_elem.getparent()
            is_deleted = False
            current = parent
            while current is not None:
                if current.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}del':
                    is_deleted = True
                    break
                current = current.getparent()

            if not is_deleted and t_elem.text:
                text_parts.append(t_elem.text)

        return ''.join(text_parts)
    
    def parse(self) -> dict:
        """
        解析 Word 文档
        
        Returns:
            包含解析结果的字典
        """
        # 提取段落文本（包含修订标记）
        self.paragraphs = []
        for para in self.document.paragraphs:
            # 使用 XML 解析获取包含修订的完整文本
            full_para_text = self.extract_text_with_revisions(para._element)
            
            if full_para_text.strip():  # 跳过空段落
                self.paragraphs.append({
                    "text": full_para_text.strip(),
                    "style": para.style.name if para.style else None,
                })
        
        # 提取表格内容
        self.tables = []
        for table_idx, table in enumerate(self.document.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            self.tables.append({
                "index": table_idx,
                "data": table_data,
            })
        
        # 合并全文
        self.full_text = "\n".join(p["text"] for p in self.paragraphs)
        
        # 添加表格文本到全文
        for table in self.tables:
            for row in table["data"]:
                self.full_text += "\n" + " ".join(row)
        
        return {
            "full_text": self.full_text,
            "paragraphs": self.paragraphs,
            "tables": self.tables,
            "paragraph_count": len(self.paragraphs),
            "table_count": len(self.tables),
        }
    
    def get_full_text(self) -> str:
        """获取完整文本"""
        if not self.full_text:
            self.parse()
        return self.full_text
    
    def get_paragraphs_text(self) -> str:
        """获取段落文本（不含表格）"""
        return "\n".join(p["text"] for p in self.paragraphs)
