"""
报告导出模块 (Report Exporter)
提供 Word/Excel/PDF/JSON/ZIP 多格式的专业级报告导出

每个导出函数接受统一的参数签名:
    comparison_result : dict   – Comparator.compare() 的输出
    summary            : dict   – Comparator.get_summary() 的输出
    word_text          : str    – Word 文档全文
    pdf_text           : str    – PDF OCR 识别全文
    diff_list          : list   – build_diff_list() 的输出（可选）
"""

import os
import io
import json
import zipfile
import textwrap
from datetime import datetime
from typing import Optional, List

# ============================================================
# 辅助函数
# ============================================================

def _normalize_diff_items(comparison_result: dict) -> dict:
    """从 comparison_result 中提取所有差异项，按类别归并"""

    def _collect(category_key: str, risk: str) -> list:
        section = comparison_result.get(category_key, {})
        items = []
        for item in section.get("missing_in_pdf", []):
            items.append({
                "raw": item.get("raw", ""),
                "keyword": item.get("keyword", ""),
                "context": item.get("context", item.get("phrase", "")),
                "normalized": item.get("normalized", ""),
                "direction": "missing_in_pdf",
                "risk": risk,
                "category": category_key,
            })
        for item in section.get("extra_in_pdf", []):
            items.append({
                "raw": item.get("raw", ""),
                "keyword": item.get("keyword", ""),
                "context": item.get("context", item.get("phrase", "")),
                "normalized": item.get("normalized", ""),
                "direction": "extra_in_pdf",
                "risk": risk,
                "category": category_key,
            })
        return items

    return {
        "amounts_digits": _collect("amounts_digits", "high"),
        "amounts_words": _collect("amounts_words", "high"),
        "dates": _collect("dates", "medium"),
        "numbers": _collect("numbers", "medium"),
        "percentages": _collect("percentages", "low"),
    }


def _risk_color_hex(risk: str) -> str:
    """风险等级 → 颜色"""
    return {"high": "D32F2F", "medium": "F57C00", "low": "1976D2"}.get(risk, "666666")


def _risk_color_rgb(risk: str) -> tuple:
    """风险等级 → RGB 元组"""
    return {
        "high": (0xD3, 0x2F, 0x2F),
        "medium": (0xF5, 0x7C, 0x00),
        "low": (0x19, 0x76, 0xD2),
    }.get(risk, (0x66, 0x66, 0x66))


def _risk_label_cn(risk: str) -> str:
    return {"high": "高风险", "medium": "中风险", "low": "低风险"}.get(risk, risk)


def _category_label_cn(category_key: str) -> str:
    _map = {
        "amounts_digits": "金额数字",
        "amounts_words": "大写金额",
        "dates": "日期",
        "numbers": "数字",
        "percentages": "百分比",
    }
    return _map.get(category_key, category_key)


def _direction_label_cn(direction: str) -> str:
    return "PDF 缺失" if direction == "missing_in_pdf" else "PDF 多出"


# ============================================================
# 1. Word Redline Export (DOCX)
# ============================================================

def export_redline_docx(
    comparison_result: dict,
    summary: dict,
    word_text: str,
    pdf_text: str,
    diff_list: Optional[List[dict]] = None,
) -> bytes:
    """
    生成带修订痕迹（Redline）的 .docx 文件

    - 摘要统计表格在文档开头
    - Word 中缺失于 PDF 的内容：红色 + 删除线
    - PDF 中多出的内容：蓝色 + 下划线
    - 正确处理中文编码
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # --- 全局默认字体（中文兼容） ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # --- 标题 ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("合同扫描件比对报告（修订版）")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # 元信息
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # 空行

    # ===== 摘要统计表 =====
    doc.add_heading("一、比对摘要", level=1)

    summary_table = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
    hdr_cells = summary_table.rows[0].cells
    headers = ["指标", "数值", "说明"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True

    risk_text = "存在金额相关差异，请重点核查！" if summary.get("has_critical_diff") else "未发现金额差异"

    rows_data = [
        ("差异总数", str(summary.get("total_diffs", 0)), ""),
        ("金额差异",
         str(summary.get("has_critical_diff", False)),
         risk_text),
    ]
    if summary.get("diff_details"):
        for d in summary["diff_details"]:
            rows_data.append((
                f"  {d['type']}",
                f"缺失 {d.get('missing', 0)} / 多出 {d.get('extra', 0)}",
                "",
            ))

    for label, val, note in rows_data:
        row = summary_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = val
        row.cells[2].text = note

    doc.add_paragraph()

    # ===== 差异明细 =====
    doc.add_heading("二、差异明细（修订标记说明）", level=1)

    # 图例
    legend_para = doc.add_paragraph()
    legend_run = legend_para.add_run("■ 标记说明：")
    legend_run.bold = True
    legend_run.font.size = Pt(9)

    legend_para2 = doc.add_paragraph()
    lr1 = legend_para2.add_run("红色删除线 = Word 原文有，PDF 扫描件缺失（可能遗漏）")
    lr1.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
    lr1.font.size = Pt(9)

    legend_para3 = doc.add_paragraph()
    lr2 = legend_para3.add_run("蓝色下划线 = PDF 扫描件多出，Word 原文无（可能 OCR 幻觉或新增）")
    lr2.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
    lr2.font.size = Pt(9)

    doc.add_paragraph()

    # ===== 带修订标记的正文 =====
    doc.add_heading("三、修订对照", level=1)
    doc.add_heading("Word 原文（红色删除线 = PDF 缺失内容）", level=2)

    normalized = _normalize_diff_items(comparison_result)

    # 收集所有缺失项文本（Word 有但 PDF 无）
    missing_texts = set()
    for cat_items in normalized.values():
        for item in cat_items:
            if item["direction"] == "missing_in_pdf" and item["raw"]:
                missing_texts.add(item["raw"])

    # 按长度降序排列以避免短串错误替换长串
    missing_sorted = sorted(missing_texts, key=len, reverse=True)

    # 逐段处理 Word 文本
    for para_text in word_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5

        if not para_text.strip():
            continue

        _render_annotated_paragraph(p, para_text, missing_sorted, mode="delete")

    doc.add_page_break()
    doc.add_heading("PDF 扫描件 OCR 文本（蓝色下划线 = PDF 多出内容）", level=2)

    # 收集所有多出项文本（PDF 有但 Word 无）
    extra_texts = set()
    for cat_items in normalized.values():
        for item in cat_items:
            if item["direction"] == "extra_in_pdf" and item["raw"]:
                extra_texts.add(item["raw"])
    extra_sorted = sorted(extra_texts, key=len, reverse=True)

    for para_text in pdf_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5

        if not para_text.strip():
            continue

        _render_annotated_paragraph(p, para_text, extra_sorted, mode="insert")

    # 写入 BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _render_annotated_paragraph(para, text: str, terms: list, mode: str) -> None:
    """
    在段落中以指定模式标注文本片段

    mode="delete": 匹配到的词 → 红色删除线；其余正常
    mode="insert": 匹配到的词 → 蓝色下划线；其余正常
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    # 使用 Unicode 私有区占位符（XML 兼容，不会与中文文本冲突）
    PH = "\ue000"
    replacements = []

    working = text
    for i, term in enumerate(terms):
        idx = working.find(term)
        if idx >= 0:
            token = f"{PH}{i}{PH}"
            replacements.append((i, term))
            working = working.replace(term, token, 1)

    # 按占位符切分
    segments = []
    last = 0
    for i, ch in enumerate(working):
        if ch == PH:
            if i > last:
                segments.append(("normal", working[last:i]))
            # 找到结束占位符
            next_ph = working.find(PH, i + 1)
            if next_ph > i:
                try:
                    idx = int(working[i + 1:next_ph])
                    for ri, rterm in replacements:
                        if ri == idx:
                            segments.append((mode, rterm))
                            break
                except (ValueError, IndexError):
                    segments.append(("normal", working[i:next_ph + 1]))
                last = next_ph + 1
            else:
                last = i
    if last < len(working):
        segments.append(("normal", working[last:]))

    for seg_type, seg_text in segments:
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)

        if seg_type == "delete":
            run.font.strike = True
            run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
        elif seg_type == "insert":
            run.underline = True
            run.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)


# ============================================================
# 2. Excel Diff Export (XLSX)
# ============================================================

def export_diff_excel(
    comparison_result: dict,
    summary: dict,
    word_text: str,
    pdf_text: str,
    diff_list: Optional[List[dict]] = None,
) -> bytes:
    """
    生成多工作表的 .xlsx 差异报告

    Sheets:
        "差异摘要"      – 总体统计
        "金额差异"      – 高风险（红色背景）
        "日期差异"      – 中风险（橙色背景）
        "数字差异"      – 中风险（橙色背景）
        "百分比差异"    – 低风险（蓝色背景）
        "全文差异"      – 原始 diff 列表
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    # ----- 通用样式 -----
    header_font = Font(name="Microsoft YaHei", bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="1A1A2E", end_color="1A1A2E", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    body_font = Font(name="Microsoft YaHei", size=10)
    body_alignment = Alignment(vertical="center", wrap_text=True)

    thin_border = Border(
        left=Side(style="thin", color="D0D5DD"),
        right=Side(style="thin", color="D0D5DD"),
        top=Side(style="thin", color="D0D5DD"),
        bottom=Side(style="thin", color="D0D5DD"),
    )

    fills = {
        "high": PatternFill(start_color="FFEBEE", end_color="FFEBEE", fill_type="solid"),
        "medium": PatternFill(start_color="FFF3E0", end_color="FFF3E0", fill_type="solid"),
        "low": PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid"),
    }
    header_fills = {
        "high": PatternFill(start_color="D32F2F", end_color="D32F2F", fill_type="solid"),
        "medium": PatternFill(start_color="F57C00", end_color="F57C00", fill_type="solid"),
        "low": PatternFill(start_color="1976D2", end_color="1976D2", fill_type="solid"),
    }

    def _write_sheet_header(ws, headers: list, col_widths: list, risk: str = None):
        fill = header_fills.get(risk, header_fill) if risk else header_fill
        for col_idx, (h, w) in enumerate(zip(headers, col_widths), 1):
            cell = ws.cell(row=1, column=col_idx, value=h)
            cell.font = header_font
            cell.fill = fill
            cell.alignment = header_alignment
            cell.border = thin_border
            ws.column_dimensions[get_column_letter(col_idx)].width = w

    def _write_row(ws, row_idx: int, values: list, risk: str = ""):
        fill = fills.get(risk, None)
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = body_font
            cell.alignment = body_alignment
            cell.border = thin_border
            if fill and risk:
                cell.fill = fill

    # ===================== Sheet 1: 差异摘要 =====================
    ws_summary = wb.active
    ws_summary.title = "差异摘要"

    summary_headers = ["指标", "数值", "说明"]
    summary_widths = [18, 14, 48]
    _write_sheet_header(ws_summary, summary_headers, summary_widths)

    summary_rows = [
        ("差异总数", summary.get("total_diffs", 0), ""),
        ("是否存在金额差异", "是" if summary.get("has_critical_diff") else "否",
         "金额差异属于高风险项，建议重点人工核查"),
        ("比对时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "自动生成"),
    ]
    if summary.get("diff_details"):
        for d in summary["diff_details"]:
            summary_rows.append((
                d["type"],
                f"缺失 {d.get('missing', 0)} / 多出 {d.get('extra', 0)}",
                "",
            ))

    risk_map_detail = {"金额数字": "high", "大写金额": "high", "日期": "medium", "数字": "medium", "百分比": "low"}
    for row_idx, (label, val, note) in enumerate(summary_rows, 2):
        r = risk_map_detail.get(label, "")
        _write_row(ws_summary, row_idx, [label, str(val), note], risk=r)

    # ===================== Sheet 2-5: 按类别分表 =====================
    sheet_configs = [
        ("amounts_digits", "金额差异", "high", ["方向", "差异值", "关键词", "上下文", "标准化值"]),
        ("amounts_words", "大写金额差异", "high", ["方向", "差异值", "备注"]),
        ("dates", "日期差异", "medium", ["方向", "日期原文", "标准化日期"]),
        ("numbers", "数字差异", "medium", ["方向", "数字", "上下文"]),
        ("percentages", "百分比差异", "low", ["方向", "差异值", "标准化值"]),
    ]

    normalized = _normalize_diff_items(comparison_result)

    for cat_key, sheet_name, risk, headers in sheet_configs:
        ws = wb.create_sheet(title=sheet_name)
        widths = [14] + [30] * (len(headers) - 1)
        _write_sheet_header(ws, headers, widths, risk=risk)

        items = normalized.get(cat_key, [])
        for row_idx, item in enumerate(items, 2):
            direction = _direction_label_cn(item["direction"])
            raw = item["raw"]
            if cat_key == "amounts_digits":
                values = [direction, raw, item.get("keyword", ""), item.get("context", ""), str(item.get("normalized", ""))]
            elif cat_key == "amounts_words":
                values = [direction, raw, ""]
            elif cat_key == "dates":
                values = [direction, raw, str(item.get("normalized", ""))]
            elif cat_key == "numbers":
                values = [direction, raw, item.get("context", "")]
            elif cat_key == "percentages":
                values = [direction, raw, str(item.get("normalized", ""))]
            else:
                values = [direction, raw]
            _write_row(ws, row_idx, values, risk=risk)

    # ===================== Sheet 6: 全文差异 =====================
    ws_full = wb.create_sheet(title="全文差异")
    full_headers = ["风险等级", "类型", "差异内容", "关键字/短语"]
    full_widths = [12, 14, 36, 30]
    _write_sheet_header(ws_full, full_headers, full_widths)

    if diff_list:
        for row_idx, d in enumerate(diff_list, 2):
            phrase = d.get("phrase", d.get("keyword", ""))
            if len(phrase) > 60:
                phrase = phrase[:57] + "..."
            values = [
                _risk_label_cn(d["risk"]),
                d.get("type", ""),
                d.get("text", ""),
                phrase,
            ]
            _write_row(ws_full, row_idx, values, risk=d["risk"])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 3. PDF Report Export (reportlab)
# ============================================================

def _register_chinese_font():
    """注册中文字体，返回字体名称"""
    import reportlab.rl_config
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "MicrosoftYaHei"
    # 防止重复注册
    try:
        pdfmetrics.getFont(font_name)
        return font_name, font_name + "Bold"
    except Exception:
        pass

    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyh.ttf",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "/System/Library/Fonts/PingFang.ttc",
    ]
    bold_paths = [
        "C:/Windows/Fonts/msyhbd.ttc",
        "C:/Windows/Fonts/msyhbd.ttf",
    ]

    font_path = None
    bold_path = None

    for p in font_paths:
        if os.path.isfile(p):
            font_path = p
            break
    for p in bold_paths:
        if os.path.isfile(p):
            bold_path = p
            break

    bold_name = font_name + "Bold"

    if font_path:
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_path, subfontIndex=0))
        except Exception:
            pdfmetrics.registerFont(TTFont(font_name, font_path))
    if bold_path and bold_path != font_path:
        try:
            pdfmetrics.registerFont(TTFont(bold_name, bold_path, subfontIndex=0))
        except Exception:
            try:
                pdfmetrics.registerFont(TTFont(bold_name, bold_path))
            except Exception:
                bold_name = font_name

    return font_name, bold_name


def export_pdf_report(
    comparison_result: dict,
    summary: dict,
    word_text: str,
    pdf_text: str,
    diff_list: Optional[List[dict]] = None,
) -> bytes:
    """
    生成专业排版的 PDF 报告（使用 reportlab）

    包含：封面页、摘要统计、差异明细表、置信度信息
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm, cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.colors import HexColor, black, white
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, KeepTogether,
    )
    from reportlab.platypus.flowables import HRFlowable

    font_name, bold_name = _register_chinese_font()

    # ----- 颜色 -----
    COLOR_DARK = HexColor("#1A1A2E")
    COLOR_RED = HexColor("#D32F2F")
    COLOR_ORANGE = HexColor("#F57C00")
    COLOR_BLUE = HexColor("#1976D2")
    COLOR_GRAY = HexColor("#999999")
    COLOR_LIGHT_BG = HexColor("#F8F9FB")

    WIDTH, HEIGHT = A4

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        title="合同比对报告",
        author="Contract Comparator",
    )

    # ----- 段落样式 -----
    styles = getSampleStyleSheet()

    style_title = ParagraphStyle(
        "CNTitle", fontName=bold_name, fontSize=22, leading=30,
        alignment=TA_CENTER, textColor=COLOR_DARK, spaceAfter=6 * mm,
    )
    style_subtitle = ParagraphStyle(
        "CNSubtitle", fontName=font_name, fontSize=10, leading=14,
        alignment=TA_CENTER, textColor=COLOR_GRAY, spaceAfter=10 * mm,
    )
    style_h1 = ParagraphStyle(
        "CNH1", fontName=bold_name, fontSize=14, leading=20,
        textColor=COLOR_DARK, spaceBefore=8 * mm, spaceAfter=4 * mm,
    )
    style_h2 = ParagraphStyle(
        "CNH2", fontName=bold_name, fontSize=12, leading=16,
        textColor=COLOR_DARK, spaceBefore=6 * mm, spaceAfter=3 * mm,
    )
    style_body = ParagraphStyle(
        "CNBody", fontName=font_name, fontSize=9, leading=14,
        textColor=black, spaceAfter=2 * mm,
    )
    style_body_small = ParagraphStyle(
        "CNBodySmall", fontName=font_name, fontSize=8, leading=11,
        textColor=COLOR_GRAY,
    )
    style_cell = ParagraphStyle(
        "CNCell", fontName=font_name, fontSize=8, leading=11,
        textColor=black,
    )
    style_cell_bold = ParagraphStyle(
        "CNCellBold", fontName=bold_name, fontSize=8, leading=11,
        textColor=black,
    )

    story = []

    # ---------- 封面 ----------
    story.append(Spacer(1, 30 * mm))
    story.append(Paragraph("合同扫描件比对报告", style_title))
    story.append(Paragraph(
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>"
        f"引擎版本：v2.0 · 本地离线",
        style_subtitle,
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_DARK))
    story.append(Spacer(1, 8 * mm))

    # ---------- 摘要 ----------
    story.append(Paragraph("一、比对摘要", style_h1))

    total = summary.get("total_diffs", 0)
    has_critical = summary.get("has_critical_diff", False)

    summary_data = [
        [
            Paragraph("<b>指标</b>", style_cell_bold),
            Paragraph("<b>数值</b>", style_cell_bold),
            Paragraph("<b>说明</b>", style_cell_bold),
        ],
        [
            Paragraph("差异总数", style_cell),
            Paragraph(str(total), style_cell),
            Paragraph("", style_cell),
        ],
    ]
    if has_critical:
        summary_data.append([
            Paragraph("金额差异", style_cell),
            Paragraph("是", style_cell),
            Paragraph('<font color="#D32F2F">⚠ 存在金额相关差异，请重点核查！</font>', style_cell),
        ])
    else:
        summary_data.append([
            Paragraph("金额差异", style_cell),
            Paragraph("否", style_cell),
            Paragraph('<font color="#2E7D32">✅ 未发现金额差异</font>', style_cell),
        ])

    if summary.get("diff_details"):
        for d in summary["diff_details"]:
            summary_data.append([
                Paragraph(f"  {d['type']}", style_cell),
                Paragraph(f"缺失 {d.get('missing', 0)} / 多出 {d.get('extra', 0)}", style_cell),
                Paragraph("", style_cell),
            ])

    summary_table = Table(summary_data, colWidths=[60 * mm, 50 * mm, 60 * mm])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), COLOR_DARK),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#D0D5DD")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [COLOR_LIGHT_BG, white]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 6 * mm))

    # ---------- 差异明细（合并表格） ----------
    story.append(Paragraph("二、差异明细", style_h1))

    normalized = _normalize_diff_items(comparison_result)
    all_diffs = []
    for cat_key, cat_label in [
        ("amounts_digits", "金额数字"),
        ("amounts_words", "大写金额"),
        ("dates", "日期"),
        ("numbers", "数字"),
        ("percentages", "百分比"),
    ]:
        for item in normalized.get(cat_key, []):
            all_diffs.append((cat_label, item["risk"], item["direction"], item["raw"], item.get("context", "")))

    if all_diffs:
        table_data = [[
            Paragraph("<b>类别</b>", style_cell_bold),
            Paragraph("<b>风险</b>", style_cell_bold),
            Paragraph("<b>方向</b>", style_cell_bold),
            Paragraph("<b>差异值</b>", style_cell_bold),
            Paragraph("<b>上下文</b>", style_cell_bold),
        ]]
        for cat_label, risk, direction, raw, ctx in all_diffs:
            color = _risk_color_hex(risk)
            table_data.append([
                Paragraph(cat_label, style_cell),
                Paragraph(f'<font color="#{color}">{_risk_label_cn(risk)}</font>', style_cell),
                Paragraph(_direction_label_cn(direction), style_cell),
                Paragraph(raw, style_cell),
                Paragraph(ctx[:60] if ctx else "", style_cell),
            ])

        col_widths = [28 * mm, 22 * mm, 24 * mm, 42 * mm, 54 * mm]
        diff_table = Table(table_data, colWidths=col_widths)
        diff_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), COLOR_DARK),
            ("TEXTCOLOR", (0, 0), (-1, 0), white),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#D0D5DD")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [COLOR_LIGHT_BG, white]),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(diff_table)
    else:
        story.append(Paragraph(
            '<font color="#2E7D32">✅ 未发现差异，所有关键字段一致。</font>',
            style_body,
        ))

    story.append(Spacer(1, 4 * mm))

    # ---------- 统计信息 ----------
    story.append(HRFlowable(width="100%", thickness=0.5, color=COLOR_GRAY))
    story.append(Paragraph(
        f"报告结束 · 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        style_body_small,
    ))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 4. JSON API Export
# ============================================================

def export_json_api(
    comparison_result: dict,
    summary: dict,
    word_text: str,
    pdf_text: str,
    diff_list: Optional[List[dict]] = None,
) -> str:
    """
    生成标准化的 JSON 格式，适合 API 集成

    Schema:
    {
      "schema_version": "2.0",
      "timestamp": "ISO8601",
      "metadata": {...},
      "summary": {...},
      "differences": [
        {
          "id": 1,
          "category": "amounts_digits",
          "category_label": "金额数字",
          "risk_level": "high",
          "direction": "missing_in_pdf",
          "direction_label": "PDF 缺失",
          "value": "...",
          "keyword": "...",
          "context": "...",
          "normalized_value": "...",
          "source_position": {...},
          "confidence": null
        }
      ],
      "full_text_diff": null,
      "statistics": {...}
    }
    """
    normalized = _normalize_diff_items(comparison_result)

    diff_records = []
    diff_id = 0
    for cat_key, items in normalized.items():
        for item in items:
            diff_id += 1
            diff_records.append({
                "id": diff_id,
                "category": cat_key,
                "category_label": _category_label_cn(cat_key),
                "risk_level": item["risk"],
                "direction": item["direction"],
                "direction_label": _direction_label_cn(item["direction"]),
                "value": item["raw"],
                "keyword": item.get("keyword", ""),
                "context": item.get("context", ""),
                "normalized_value": str(item.get("normalized", "")),
                "source_position": {
                    "word_line": None,
                    "pdf_page": None,
                    "pdf_line": None,
                },
                "confidence": None,
            })

    total_by_risk = {"high": 0, "medium": 0, "low": 0}
    for d in diff_records:
        total_by_risk[d["risk_level"]] += 1

    payload = {
        "schema_version": "2.0",
        "timestamp": datetime.now().isoformat(),
        "metadata": {
            "engine": "contract_comparator",
            "engine_version": "2.0",
            "generated_by": "report_exporter.export_json_api",
        },
        "summary": {
            "total_differences": summary.get("total_diffs", 0),
            "has_critical_diff": summary.get("has_critical_diff", False),
            "diff_details": summary.get("diff_details", []),
        },
        "differences": diff_records,
        "full_text_diff": None,
        "statistics": {
            "total_differences": len(diff_records),
            "by_risk_level": total_by_risk,
            "by_category": {
                cat: sum(1 for d in diff_records if d["category"] == cat)
                for cat in set(d["category"] for d in diff_records)
            },
            "by_direction": {
                "missing_in_pdf": sum(1 for d in diff_records if d["direction"] == "missing_in_pdf"),
                "extra_in_pdf": sum(1 for d in diff_records if d["direction"] == "extra_in_pdf"),
            },
        },
    }

    return json.dumps(payload, ensure_ascii=False, indent=2)


# ============================================================
# 5. Comparison Report ZIP
# ============================================================

def export_full_package(
    comparison_result: dict,
    summary: dict,
    word_text: str,
    pdf_text: str,
    diff_list: Optional[List[dict]] = None,
) -> bytes:
    """
    生成包含所有报告格式 + 原始文本的 ZIP 包

    目录结构:
        reports/
        ├── 合同比对报告_修订版.docx
        ├── 合同比对报告_差异明细.xlsx
        ├── 合同比对报告.pdf
        ├── 合同比对报告.json
        ├── 原始文本/
        │   ├── word_原文.txt
        │   └── pdf_ocr文本.txt
    """
    buf = io.BytesIO()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # --- Word Redline ---
        try:
            docx_bytes = export_redline_docx(comparison_result, summary, word_text, pdf_text, diff_list)
            zf.writestr(f"reports/{timestamp}/合同比对报告_修订版.docx", docx_bytes)
        except Exception:
            pass  # 静默跳过单个格式失败，不中断 ZIP 打包

        # --- Excel ---
        try:
            xlsx_bytes = export_diff_excel(comparison_result, summary, word_text, pdf_text, diff_list)
            zf.writestr(f"reports/{timestamp}/合同比对报告_差异明细.xlsx", xlsx_bytes)
        except Exception:
            pass

        # --- PDF ---
        try:
            pdf_bytes = export_pdf_report(comparison_result, summary, word_text, pdf_text, diff_list)
            zf.writestr(f"reports/{timestamp}/合同比对报告.pdf", pdf_bytes)
        except Exception:
            pass

        # --- JSON API ---
        try:
            json_str = export_json_api(comparison_result, summary, word_text, pdf_text, diff_list)
            zf.writestr(f"reports/{timestamp}/合同比对报告.json", json_str.encode("utf-8"))
        except Exception:
            pass

        # --- 原始文本 ---
        zf.writestr(
            f"reports/{timestamp}/原始文本/word_原文.txt",
            word_text.encode("utf-8"),
        )
        zf.writestr(
            f"reports/{timestamp}/原始文本/pdf_ocr文本.txt",
            pdf_text.encode("utf-8"),
        )

    buf.seek(0)
    return buf.getvalue()